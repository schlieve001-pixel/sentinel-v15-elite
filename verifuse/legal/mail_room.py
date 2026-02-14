"""
VeriFuse Mail Room — Attorney Solicitation Letter Generator
============================================================
Colorado attorneys cannot cold-call surplus fund owners (CO Rule 7.3).
They MUST send written correspondence. This module generates
compliant .docx solicitation letters from pipeline WHALE/PRIME leads.

COMPLIANCE:
  Every letter includes the mandatory footer:
  "ADVERTISING MATERIAL - COMPLIANT WITH CO RULE 7.3"

USAGE:
  from verifuse.legal.mail_room import generate_letter, run_mail_room
  generate_letter(asset_dict, attorney_info, "output.docx")
  run_mail_room()  # Batch: all ATTORNEY-class Whales/Primes
"""

import os
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ============================================================================
# CONSTANTS
# ============================================================================

DB_PATH = Path(__file__).resolve().parent.parent / "data" / "verifuse.db"

COMPLIANCE_FOOTER = "ADVERTISING MATERIAL - COMPLIANT WITH CO RULE 7.3"

DEFAULT_ATTORNEY = {
    "firm_name": "[LAW FIRM NAME]",
    "attorney_name": "[ATTORNEY NAME, ESQ.]",
    "bar_number": "[CO BAR #]",
    "address_line1": "[FIRM ADDRESS LINE 1]",
    "address_line2": "[CITY, STATE ZIP]",
    "phone": "[PHONE NUMBER]",
    "email": "[EMAIL]",
}

BLACK = RGBColor(0x0A, 0x19, 0x2F)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_GREEN = RGBColor(0x00, 0x80, 0x3E)


# ============================================================================
# HELPERS
# ============================================================================

def _fmt_money(amount) -> str:
    """Format a number as $X,XXX.XX."""
    if amount is None:
        return "$[AMOUNT]"
    try:
        return f"${float(amount):,.2f}"
    except (ValueError, TypeError):
        return "$[AMOUNT]"


def _add_paragraph(doc, text, size=11, bold=False, color=None,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                   font_name="Calibri"):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


# ============================================================================
# LETTER GENERATOR
# ============================================================================

def generate_letter(
    asset: Dict,
    attorney: Optional[Dict] = None,
    output_path: Optional[str] = None,
) -> str:
    """
    Generate a single .docx solicitation letter for a surplus fund lead.

    The letter follows the required format:
      1. Firm letterhead
      2. Date + recipient address
      3. RE: line with surplus amount and case number (THE HOOK)
      4. Body explaining unclaimed funds and attorney services
      5. Call to action
      6. Compliance footer: "ADVERTISING MATERIAL - COMPLIANT WITH CO RULE 7.3"

    Args:
        asset:       Asset dict from pipeline or hunter engine.
        attorney:    Dict with firm_name, attorney_name, bar_number, etc.
        output_path: Where to save. Auto-generated if None.

    Returns:
        Path to the generated .docx file.
    """
    atty = attorney or DEFAULT_ATTORNEY
    doc = Document()

    # ---- PAGE SETUP ----
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ---- FIRM LETTERHEAD ----
    _add_paragraph(doc, atty["firm_name"], size=14, bold=True, color=BLACK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_paragraph(doc, atty.get("address_line1", ""), size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _add_paragraph(doc, atty.get("address_line2", ""), size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    contact_line = f"Tel: {atty.get('phone', '')}  |  {atty.get('email', '')}"
    _add_paragraph(doc, contact_line, size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ---- DIVIDER ----
    _add_paragraph(doc, "_" * 72, size=8, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ---- DATE ----
    today = datetime.now().strftime("%B %d, %Y")
    _add_paragraph(doc, today, size=11, color=DARK_GRAY, space_after=12)

    # ---- RECIPIENT ----
    owner = asset.get("owner_of_record", "[PROPERTY OWNER]")
    prop_addr = asset.get("property_address", "[PROPERTY ADDRESS]")
    mail_addr = asset.get("_mailing_address") or prop_addr

    _add_paragraph(doc, owner, size=11, bold=True, color=DARK_GRAY, space_after=1)
    _add_paragraph(doc, mail_addr, size=11, color=DARK_GRAY, space_after=12)

    # ---- THE HOOK: RE: LINE ----
    surplus = asset.get("estimated_surplus")
    case_num = asset.get("case_number", "[CASE NUMBER]")
    county = asset.get("county", "[COUNTY]")

    re_line = f"RE: Unclaimed Funds of {_fmt_money(surplus)} - Case #{case_num}"
    _add_paragraph(doc, re_line, size=12, bold=True, color=BLACK, space_after=12)

    # ---- GREETING ----
    _add_paragraph(doc, f"Dear {owner},", size=11, color=DARK_GRAY, space_after=8)

    # ---- BODY PARAGRAPH 1: The Alert ----
    body1 = (
        f"Our office has identified unclaimed surplus funds in the amount of "
        f"{_fmt_money(surplus)} that may belong to you. These funds resulted "
        f"from a foreclosure sale involving real property located at "
        f"{prop_addr}, in {county} County, Colorado."
    )
    _add_paragraph(doc, body1, size=11, color=DARK_GRAY, space_after=8)

    # ---- BODY PARAGRAPH 2: The Urgency ----
    sale_date = asset.get("sale_date", "")
    days_remaining = asset.get("days_remaining")
    asset_type = str(asset.get("asset_type", "")).upper()
    if asset_type in ("TAX_OVERPAYMENT", "TAX_DEED_SURPLUS"):
        statute = "C.R.S. Section 39-11-151"
    else:
        statute = "C.R.S. Section 38-38-111"

    body2 = (
        f"Under {statute}, these funds are held by the {county} County "
        f"Public Trustee and are available for claim by the rightful owner. "
        f"However, there is a statutory deadline to file your claim. "
        f"After this deadline, the funds may be transferred to the Colorado "
        f"State Treasury, making recovery significantly more difficult."
    )
    _add_paragraph(doc, body2, size=11, color=DARK_GRAY, space_after=8)

    # ---- BODY PARAGRAPH 3: The Offer ----
    body3 = (
        f"Our firm specializes in recovering surplus funds from foreclosure "
        f"sales in Colorado. We handle the entire legal process, including "
        f"filing the necessary petition with the court, locating and "
        f"assembling all required documentation, and representing your "
        f"interests through to the disbursement of funds."
    )
    _add_paragraph(doc, body3, size=11, color=DARK_GRAY, space_after=8)

    # ---- BODY PARAGRAPH 4: No Upfront Cost ----
    body4 = (
        "There is no upfront cost to you. Our fee is contingent upon "
        "successful recovery of your funds. If we do not recover any "
        "money, you owe us nothing."
    )
    _add_paragraph(doc, body4, size=11, bold=True, color=DARK_GRAY, space_after=8)

    # ---- CALL TO ACTION ----
    body5 = (
        "To discuss this matter, please contact our office at your earliest "
        f"convenience by calling {atty.get('phone', '[PHONE]')} or by "
        f"emailing {atty.get('email', '[EMAIL]')}. Time is of the essence."
    )
    _add_paragraph(doc, body5, size=11, color=DARK_GRAY, space_after=16)

    # ---- SIGNATURE ----
    _add_paragraph(doc, "Sincerely,", size=11, color=DARK_GRAY, space_after=24)
    _add_paragraph(doc, atty["attorney_name"], size=11, bold=True,
                   color=DARK_GRAY, space_after=1)
    _add_paragraph(doc, atty["firm_name"], size=10, color=MID_GRAY, space_after=1)
    bar = atty.get("bar_number", "")
    if bar:
        _add_paragraph(doc, f"Colorado Bar #{bar}", size=9, color=MID_GRAY,
                       space_after=1)

    # ---- COMPLIANCE FOOTER (MANDATORY) ----
    _add_paragraph(doc, "", size=6, space_after=20)  # spacer
    _add_paragraph(doc, "_" * 72, size=8, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_paragraph(
        doc, COMPLIANCE_FOOTER,
        size=9, bold=True, color=RGBColor(0xCC, 0x00, 0x00),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_paragraph(
        doc,
        "This letter is being sent to you because public records indicate "
        "you may have a legal right to surplus funds. You are under no "
        "obligation to respond. You may wish to consult with another "
        "attorney of your choosing.",
        size=8, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )

    # ---- SAVE ----
    if not output_path:
        asset_id = asset.get("asset_id", "UNKNOWN")
        timestamp = datetime.now().strftime("%Y%m%d")
        output_path = f"VF_LETTER_{asset_id}_{timestamp}.docx"

    doc.save(output_path)
    return output_path


# ============================================================================
# BATCH MODE — Generate letters for all ATTORNEY Whales/Primes
# ============================================================================

def get_whale_leads(db_path: Optional[str] = None) -> List[Dict]:
    """
    Pull all ATTORNEY-class assets from the pipeline that are WHALE or PRIME
    candidates (surplus >= $25K).

    Returns list of asset dicts ready for letter generation.
    """
    path = db_path or str(DB_PATH)
    if not os.path.exists(path):
        print(f"Database not found: {path}")
        return []

    conn = sqlite3.connect(path)
    conn.row_factory = sqlite3.Row

    rows = conn.execute("""
        SELECT a.*, ls.record_class, ls.data_grade, ls.days_remaining
        FROM assets a
        JOIN legal_status ls ON a.asset_id = ls.asset_id
        WHERE ls.record_class = 'ATTORNEY'
          AND ls.data_grade IN ('GOLD', 'SILVER')
          AND ls.days_remaining > 0
          AND a.estimated_surplus >= 25000
          AND a.owner_of_record IS NOT NULL
          AND a.owner_of_record != 'Unknown'
        ORDER BY a.estimated_surplus DESC
    """).fetchall()

    leads = [dict(row) for row in rows]
    conn.close()

    print(f"Found {len(leads)} WHALE/PRIME leads with surplus >= $25K")
    return leads


def run_mail_room(
    attorney: Optional[Dict] = None,
    output_dir: str = "mail_room_output",
    db_path: Optional[str] = None,
    min_surplus: float = 25000,
) -> List[str]:
    """
    Batch generate solicitation letters for all qualifying leads.

    Args:
        attorney:    Attorney info dict. Uses placeholders if None.
        output_dir:  Directory for generated letters.
        db_path:     Path to verifuse DB.
        min_surplus: Minimum surplus to qualify. Default: $25,000.

    Returns:
        List of generated file paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    leads = get_whale_leads(db_path)

    if not leads:
        print("No qualifying leads found.")
        return []

    # Filter by minimum surplus
    leads = [l for l in leads if (l.get("estimated_surplus") or 0) >= min_surplus]

    print(f"\n{'='*70}")
    print(f"VERIFUSE MAIL ROOM")
    print(f"{'='*70}")
    print(f"Generating letters for {len(leads)} leads...")
    print(f"Minimum surplus: {_fmt_money(min_surplus)}")
    print(f"Output directory: {output_dir}/")
    print(f"{'='*70}")

    paths = []
    total_surplus = 0

    for i, lead in enumerate(leads):
        county = lead.get("county", "UNK").replace(" ", "_")
        asset_id = lead.get("asset_id", f"LEAD_{i:04d}")
        surplus = lead.get("estimated_surplus", 0)
        total_surplus += surplus

        filename = f"VF_LETTER_{county}_{asset_id}.docx"
        filepath = os.path.join(output_dir, filename)

        try:
            generate_letter(lead, attorney, filepath)
            paths.append(filepath)
            owner = lead.get("owner_of_record", "Unknown")[:30]
            print(f"  [{i+1:3d}] {county:12s} | {_fmt_money(surplus):>14s} | {owner}")
        except Exception as e:
            print(f"  [ERR] {asset_id}: {e}")

    print(f"\n{'='*70}")
    print(f"MAIL ROOM COMPLETE")
    print(f"{'='*70}")
    print(f"Letters generated:  {len(paths)}/{len(leads)}")
    print(f"Total surplus targeted: {_fmt_money(total_surplus)}")
    print(f"Output: {os.path.abspath(output_dir)}/")
    print(f"\nREMINDER: Replace [BRACKETED] placeholders with actual attorney info.")
    print(f"COMPLIANCE: All letters include CO Rule 7.3 advertising footer.")

    return paths


# ============================================================================
# CLI
# ============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--sample":
        # Generate a sample letter
        from verifuse.attorney.dossier_generator import SAMPLE_ASSET
        path = generate_letter(SAMPLE_ASSET)
        print(f"Sample letter: {path}")
    else:
        run_mail_room()
