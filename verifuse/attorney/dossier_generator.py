"""
VeriFuse Pre-Litigation Evidence Packet — DOCX Dossier Generator
=================================================================
Generates a 3-page Word Document "Forensic Report" for each asset.

PAGE 1: Executive Intelligence Summary
  - Confidential header bar
  - Liquidity/Statute/Owner status table (the "green box")
  - Google Static Map placeholder

PAGE 2: Chain of Title Audit
  - Foreclosure date, bid amount, deed link
  - The Math: (Bid) - (Debt) = (Surplus) shown explicitly
  - Data provenance table

PAGE 3: Ready-to-File Motion
  - Pre-written Motion for Disbursement of Funds
  - Placeholders filled with asset data
  - C.R.S. § 38-38-111 citation

USAGE:
  pip install python-docx
  python dossier_generator.py              # generates sample dossier
  python -c "from verifuse.attorney.dossier_generator import generate_dossier; ..."

DESIGN DECISIONS:
  - python-docx chosen over WeasyPrint for this deliverable because attorneys
    need to EDIT the motion before filing. Word is the standard.
  - The existing case_packet.py (HTML/PDF) remains for the read-only evidence packet.
  - This dossier is the EDITABLE counterpart — the attorney's working document.
"""

import os
from datetime import datetime, timedelta
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================================
# COLOR PALETTE — "Bloomberg Terminal" aesthetic in print form
# ============================================================================

BLACK = RGBColor(0x0A, 0x19, 0x2F)       # Dark navy (background headers)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_GREEN = RGBColor(0x00, 0x80, 0x3E) # Surplus green (money color)
ACCENT_RED = RGBColor(0xCC, 0x00, 0x00)   # Warning red
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)    # Body text
MID_GRAY = RGBColor(0x66, 0x66, 0x66)     # Secondary text
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)  # Table stripe


# ============================================================================
# HELPERS
# ============================================================================

def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set individual cell borders. kwargs: top, bottom, left, right with values like '4' (eighths of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val}" '
            f'w:space="0" w:color="0A192F"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_styled_paragraph(doc, text: str, size: int = 11,
                          bold: bool = False, italic: bool = False,
                          color: RGBColor = DARK_GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_after: int = 6,
                          space_before: int = 0,
                          font_name: str = "Calibri") -> Any:
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font_name
    return p


def _format_currency(amount) -> str:
    """Format a number as USD currency."""
    if amount is None:
        return "NOT AVAILABLE"
    try:
        val = float(amount)
        if val < 0:
            return f"-${abs(val):,.2f}"
        return f"${val:,.2f}"
    except (ValueError, TypeError):
        return str(amount)


def _compute_statute_status(sale_date_str: Optional[str],
                            statute_years: int = 5) -> Dict:
    """
    Compute statute status from sale date.
    Returns dict with status label, days_since_sale, days_remaining, fee_cap.
    """
    if not sale_date_str:
        return {
            "label": "INSUFFICIENT DATA",
            "days_since_sale": None,
            "days_remaining": None,
            "fee_cap": None,
            "fee_cap_label": "Unknown",
        }

    try:
        sale_date = datetime.strptime(str(sale_date_str)[:10], "%Y-%m-%d")
    except ValueError:
        return {
            "label": "DATE PARSE ERROR",
            "days_since_sale": None,
            "days_remaining": None,
            "fee_cap": None,
            "fee_cap_label": "Unknown",
        }

    days_since = (datetime.utcnow() - sale_date).days
    total_window = statute_years * 365
    days_remaining = total_window - days_since

    if days_since <= 180:
        label = "ATTORNEY EXCLUSIVE — DAY 1 ACCESS"
        fee_cap = 0.33
        fee_cap_label = "33% (No statutory cap — unregulated period)"
    elif days_since <= 730:
        label = "FINDER ELIGIBLE (C.R.S. 38-38-111 — 20% Cap)"
        fee_cap = 0.20
        fee_cap_label = "20% (C.R.S. 38-38-111)"
    elif days_since <= total_window:
        label = "STATE TREASURY (RUUPA — 10% Cap)"
        fee_cap = 0.10
        fee_cap_label = "10% (Revised Uniform Unclaimed Property Act)"
    else:
        label = "ESCHEATMENT RISK — WINDOW CLOSING"
        fee_cap = 0.10
        fee_cap_label = "10% or case-by-case"

    return {
        "label": label,
        "days_since_sale": days_since,
        "days_remaining": max(days_remaining, 0),
        "fee_cap": fee_cap,
        "fee_cap_label": fee_cap_label,
    }


# ============================================================================
# PAGE 1: EXECUTIVE INTELLIGENCE SUMMARY
# ============================================================================

def _build_page_1(doc: Document, data: Dict):
    """
    PAGE 1: Executive Intelligence Summary

    Layout:
      [HEADER BAR — dark navy, white text]
      [Date / Asset ID / County metadata]
      [THE BOX — Key metrics table with green surplus]
      [OWNER INTELLIGENCE — Absentee analysis]
      [MAP PLACEHOLDER — Google Static Maps]
    """

    # ---- HEADER BAR ----
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    _set_cell_shading(header_cell, "0A192F")

    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(12)
    hp.paragraph_format.space_after = Pt(12)

    run = hp.add_run("VERIFUSE FORENSIC REPORT")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = "Calibri"

    run2 = hp.add_run("  |  CONFIDENTIAL")
    run2.font.size = Pt(12)
    run2.font.bold = False
    run2.font.color.rgb = RGBColor(0x64, 0xFF, 0xDA)  # Accent teal
    run2.font.name = "Calibri"

    # Set table width to full page
    for cell in header_table.rows[0].cells:
        cell.width = Inches(7.0)

    # ---- METADATA LINE ----
    generated = datetime.utcnow().strftime("%B %d, %Y at %H:%M UTC")
    _add_styled_paragraph(
        doc,
        f"Generated: {generated}  |  Asset: {data.get('asset_id', 'N/A')}  |  "
        f"County: {data.get('county', 'N/A')}, {data.get('state', 'CO')}",
        size=9, color=MID_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=16,
    )

    # ---- THE BOX: KEY METRICS TABLE ----
    _add_styled_paragraph(
        doc, "EXECUTIVE INTELLIGENCE SUMMARY",
        size=13, bold=True, color=BLACK, space_before=4, space_after=8,
    )

    surplus = data.get("estimated_surplus")
    statute = _compute_statute_status(
        data.get("sale_date"),
        data.get("statute_years", 5),
    )

    # Absentee detection
    is_absentee = data.get("_is_absentee", False)
    if is_absentee:
        owner_status = "ABSENTEE OWNER — Owner appears to have vacated the property"
        owner_color = "00803E"  # Green — favorable for attorney
    else:
        prop = data.get("property_address", "")
        mail = data.get("_mailing_address", "")
        if prop and mail and prop != mail:
            owner_status = "ABSENTEE OWNER — Mailing address differs from property"
            owner_color = "00803E"
        elif not mail:
            owner_status = "OWNER STATUS UNKNOWN — No mailing address on record"
            owner_color = "666666"
        else:
            owner_status = "OWNER-OCCUPANT — Owner may still reside at property"
            owner_color = "333333"

    # Build the metrics table (3 rows: Liquidity, Statute, Owner)
    metrics = doc.add_table(rows=3, cols=2)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.style = "Table Grid"

    # Set column widths
    for row in metrics.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.8)

    # Row 1: LIQUIDITY AVAILABLE
    label_cell = metrics.cell(0, 0)
    _set_cell_shading(label_cell, "0A192F")
    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after = Pt(8)
    lr = lp.add_run("LIQUIDITY\nAVAILABLE")
    lr.font.size = Pt(10)
    lr.font.bold = True
    lr.font.color.rgb = WHITE
    lr.font.name = "Calibri"

    value_cell = metrics.cell(0, 1)
    vp = value_cell.paragraphs[0]
    vp.paragraph_format.space_before = Pt(6)
    vp.paragraph_format.space_after = Pt(6)
    vr = vp.add_run(_format_currency(surplus))
    vr.font.size = Pt(22)
    vr.font.bold = True
    vr.font.color.rgb = ACCENT_GREEN if surplus and surplus > 0 else ACCENT_RED
    vr.font.name = "Calibri"

    # Fee estimate sub-line
    if surplus and statute["fee_cap"]:
        fee_est = surplus * statute["fee_cap"]
        fee_run = vp.add_run(f"\nEstimated Attorney Fee: {_format_currency(fee_est)}")
        fee_run.font.size = Pt(9)
        fee_run.font.color.rgb = MID_GRAY
        fee_run.font.name = "Calibri"

    # Row 2: STATUTE WINDOW
    label_cell_2 = metrics.cell(1, 0)
    _set_cell_shading(label_cell_2, "0A192F")
    lp2 = label_cell_2.paragraphs[0]
    lp2.paragraph_format.space_before = Pt(8)
    lp2.paragraph_format.space_after = Pt(8)
    lr2 = lp2.add_run("STATUTE\nWINDOW")
    lr2.font.size = Pt(10)
    lr2.font.bold = True
    lr2.font.color.rgb = WHITE
    lr2.font.name = "Calibri"

    value_cell_2 = metrics.cell(1, 1)
    vp2 = value_cell_2.paragraphs[0]
    vp2.paragraph_format.space_before = Pt(6)
    vp2.paragraph_format.space_after = Pt(6)
    vr2 = vp2.add_run(statute["label"])
    vr2.font.size = Pt(12)
    vr2.font.bold = True
    vr2.font.name = "Calibri"

    # Statute is green if exclusive, yellow if finder, red if treasury/escheat
    if "EXCLUSIVE" in statute["label"]:
        vr2.font.color.rgb = ACCENT_GREEN
    elif "FINDER" in statute["label"]:
        vr2.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)  # Amber
    else:
        vr2.font.color.rgb = ACCENT_RED

    # Days remaining sub-line
    if statute["days_remaining"] is not None:
        dr_text = f"\n{statute['days_remaining']} days remaining  |  Fee cap: {statute['fee_cap_label']}"
        dr_run = vp2.add_run(dr_text)
        dr_run.font.size = Pt(9)
        dr_run.font.color.rgb = MID_GRAY
        dr_run.font.name = "Calibri"

    # Row 3: OWNER STATUS
    label_cell_3 = metrics.cell(2, 0)
    _set_cell_shading(label_cell_3, "0A192F")
    lp3 = label_cell_3.paragraphs[0]
    lp3.paragraph_format.space_before = Pt(8)
    lp3.paragraph_format.space_after = Pt(8)
    lr3 = lp3.add_run("OWNER\nSTATUS")
    lr3.font.size = Pt(10)
    lr3.font.bold = True
    lr3.font.color.rgb = WHITE
    lr3.font.name = "Calibri"

    value_cell_3 = metrics.cell(2, 1)
    vp3 = value_cell_3.paragraphs[0]
    vp3.paragraph_format.space_before = Pt(6)
    vp3.paragraph_format.space_after = Pt(6)
    vr3 = vp3.add_run(owner_status)
    vr3.font.size = Pt(11)
    vr3.font.bold = True
    vr3.font.name = "Calibri"
    vr3.font.color.rgb = RGBColor(
        int(owner_color[:2], 16),
        int(owner_color[2:4], 16),
        int(owner_color[4:6], 16),
    )

    # ---- PROPERTY DETAILS ----
    _add_styled_paragraph(doc, "", size=6, space_after=2)  # Spacer

    details_table = doc.add_table(rows=4, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.style = "Table Grid"

    detail_rows = [
        ("Property Address", data.get("property_address", "NOT AVAILABLE")),
        ("Owner of Record", data.get("owner_of_record", "NOT AVAILABLE")),
        ("Case / File Number", data.get("case_number", "NOT AVAILABLE")),
        ("Asset Classification", data.get("_classification", "N/A")),
    ]
    for i, (label, value) in enumerate(detail_rows):
        lc = details_table.cell(i, 0)
        vc = details_table.cell(i, 1)
        lc.width = Inches(2.2)
        vc.width = Inches(4.8)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(9)
        lr.font.bold = True
        lr.font.color.rgb = DARK_GRAY
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(10)
        vr.font.color.rgb = DARK_GRAY
        vr.font.name = "Calibri"

    # ---- MAP PLACEHOLDER ----
    _add_styled_paragraph(
        doc, "SATELLITE INTELLIGENCE",
        size=11, bold=True, color=BLACK, space_before=16, space_after=4,
    )

    map_url = data.get("satellite_map_url", "")
    if map_url:
        _add_styled_paragraph(
            doc, f"Map URL: {map_url}",
            size=8, color=MID_GRAY, italic=True, space_after=2,
        )

    # Placeholder box for the map image
    map_table = doc.add_table(rows=1, cols=1)
    map_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    map_cell = map_table.cell(0, 0)
    _set_cell_shading(map_cell, "E8E8E8")
    map_cell.width = Inches(5.5)

    mp = map_cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(30)
    mp.paragraph_format.space_after = Pt(30)
    mr = mp.add_run("[GOOGLE STATIC MAP — INSERT IMAGE HERE]")
    mr.font.size = Pt(10)
    mr.font.color.rgb = MID_GRAY
    mr.font.italic = True
    mr.font.name = "Calibri"

    addr = data.get("property_address", "")
    if addr:
        mr2 = mp.add_run(f"\nProperty: {addr}")
        mr2.font.size = Pt(8)
        mr2.font.color.rgb = MID_GRAY
        mr2.font.name = "Calibri"


# ============================================================================
# PAGE 2: CHAIN OF TITLE AUDIT
# ============================================================================

def _build_page_2(doc: Document, data: Dict):
    """
    PAGE 2: Chain of Title Audit

    Layout:
      [HEADER]
      [Foreclosure event table]
      [THE MATH — Bid - Debt = Surplus, shown step by step]
      [Recorder / Deed link]
      [Data provenance]
    """

    # Page break
    doc.add_page_break()

    # ---- HEADER ----
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    _set_cell_shading(header_cell, "0A192F")
    header_cell.width = Inches(7.0)

    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after = Pt(8)
    run = hp.add_run("CHAIN OF TITLE AUDIT")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = "Calibri"

    # ---- FORECLOSURE EVENT TABLE ----
    _add_styled_paragraph(
        doc, "Foreclosure / Sale Event",
        size=12, bold=True, color=BLACK, space_before=14, space_after=6,
    )

    event_data = [
        ("Foreclosure / Sale Date", data.get("sale_date", "NOT AVAILABLE")),
        ("Lien Type", data.get("lien_type", "Deed of Trust")),
        ("County", f"{data.get('county', 'N/A')}, {data.get('state', 'CO')}"),
        ("Case Number", data.get("case_number", "NOT AVAILABLE")),
        ("Redemption Deadline", data.get("redemption_date", "N/A or expired")),
    ]

    event_table = doc.add_table(rows=len(event_data), cols=2)
    event_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    event_table.style = "Table Grid"

    for i, (label, value) in enumerate(event_data):
        lc = event_table.cell(i, 0)
        vc = event_table.cell(i, 1)
        lc.width = Inches(2.5)
        vc.width = Inches(4.5)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(10)
        lr.font.bold = True
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(10)
        vr.font.name = "Calibri"

    # ---- THE MATH ----
    _add_styled_paragraph(
        doc, "THE MATH — Surplus Calculation",
        size=12, bold=True, color=BLACK, space_before=20, space_after=4,
    )

    _add_styled_paragraph(
        doc,
        "Attorneys need to see the math. Here it is.",
        size=9, italic=True, color=MID_GRAY, space_after=8,
    )

    bid = data.get("overbid_amount") or data.get("bid_amount")
    debt = data.get("total_indebtedness")
    surplus = data.get("estimated_surplus")

    # If surplus wasn't provided but we have bid and debt, compute it
    if surplus is None and bid is not None and debt is not None:
        try:
            surplus = float(bid) - float(debt)
        except (ValueError, TypeError):
            pass

    math_table = doc.add_table(rows=4, cols=2)
    math_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    math_table.style = "Table Grid"

    math_rows = [
        ("Winning Bid / Sale Price", _format_currency(bid)),
        ("Total Indebtedness / Judgment", f"– {_format_currency(debt)}"),
        ("", ""),  # Divider row
        ("SURPLUS (Liquidity Available)", _format_currency(surplus)),
    ]

    for i, (label, value) in enumerate(math_rows):
        lc = math_table.cell(i, 0)
        vc = math_table.cell(i, 1)
        lc.width = Inches(3.5)
        vc.width = Inches(3.5)

        if i == 2:
            # Divider row — just a line
            _set_cell_shading(lc, "0A192F")
            _set_cell_shading(vc, "0A192F")
            lc.paragraphs[0].paragraph_format.space_before = Pt(1)
            lc.paragraphs[0].paragraph_format.space_after = Pt(1)
            vc.paragraphs[0].paragraph_format.space_before = Pt(1)
            vc.paragraphs[0].paragraph_format.space_after = Pt(1)
            continue

        if i == 3:
            # Result row — bold green
            _set_cell_shading(lc, "E8F5E9")
            _set_cell_shading(vc, "E8F5E9")

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.font.size = Pt(11)
        lr.font.bold = (i == 3)
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.font.size = Pt(14 if i == 3 else 11)
        vr.font.bold = (i == 3)
        vr.font.name = "Calibri"
        if i == 3:
            vr.font.color.rgb = ACCENT_GREEN if surplus and surplus > 0 else ACCENT_RED

    # ---- RECORDER / DEED LINK ----
    _add_styled_paragraph(
        doc, "DEED & RECORDER REFERENCE",
        size=12, bold=True, color=BLACK, space_before=20, space_after=6,
    )

    recorder_link = data.get("recorder_link", "NOT AVAILABLE — search county recorder manually")
    doc_link = data.get("document_link", "")

    ref_data = [
        ("County Recorder Search", recorder_link),
        ("Bid Sheet / Document", doc_link or "No document link on file"),
    ]

    ref_table = doc.add_table(rows=len(ref_data), cols=2)
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_table.style = "Table Grid"

    for i, (label, value) in enumerate(ref_data):
        lc = ref_table.cell(i, 0)
        vc = ref_table.cell(i, 1)
        lc.width = Inches(2.5)
        vc.width = Inches(4.5)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(9)
        lr.font.bold = True
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(8)
        vr.font.color.rgb = RGBColor(0x00, 0x66, 0xCC) if "http" in str(value) else DARK_GRAY
        vr.font.name = "Calibri"

    _add_styled_paragraph(
        doc,
        "Note: Recorder links are search URLs. Use the owner name and case number "
        "to locate the relevant filing. All data should be independently verified.",
        size=8, italic=True, color=MID_GRAY, space_before=6, space_after=4,
    )

    # ---- DATA PROVENANCE ----
    _add_styled_paragraph(
        doc, "DATA PROVENANCE",
        size=11, bold=True, color=BLACK, space_before=16, space_after=4,
    )

    prov_data = [
        ("Data Source", f"Public records — {data.get('county', 'N/A')} County, {data.get('state', 'CO')}"),
        ("Collection Date", str(data.get("created_at", "Unknown"))[:10]),
        ("Last Updated", str(data.get("updated_at", datetime.utcnow().isoformat()))[:10]),
        ("Source", data.get("source_file", "VeriFuse automated collection")),
        ("Data Grade", data.get("data_grade", "N/A")),
        ("Litigation Quality", data.get("_litigation_quality", "N/A")),
    ]

    prov_table = doc.add_table(rows=len(prov_data), cols=2)
    prov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prov_table.style = "Table Grid"

    for i, (label, value) in enumerate(prov_data):
        lc = prov_table.cell(i, 0)
        vc = prov_table.cell(i, 1)
        lc.width = Inches(2.0)
        vc.width = Inches(5.0)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(8)
        vr.font.name = "Calibri"


# ============================================================================
# PAGE 3: READY-TO-FILE MOTION
# ============================================================================

def _build_page_3(doc: Document, data: Dict):
    """
    PAGE 3: Pre-written Motion for Disbursement of Surplus Funds.

    This is a TEMPLATE — the attorney edits it before filing.
    All placeholders are filled from asset data.
    Cites C.R.S. § 38-38-111 for Colorado foreclosure surplus.
    """

    doc.add_page_break()

    # ---- COURT HEADER ----
    county = data.get("county", "[COUNTY]")
    state = data.get("state", "CO")
    case_number = data.get("case_number", "[CASE NUMBER]")
    owner = data.get("owner_of_record", "[OWNER NAME]")
    property_address = data.get("property_address", "[PROPERTY ADDRESS]")
    surplus = data.get("estimated_surplus")
    surplus_str = _format_currency(surplus)
    sale_date = data.get("sale_date", "[SALE DATE]")
    lien_type = data.get("lien_type", "Deed of Trust")

    # Determine statute citation based on asset type
    asset_type = data.get("asset_type", "FORECLOSURE_SURPLUS")
    if "TAX" in str(asset_type).upper():
        statute_cite = "C.R.S. § 39-11-151"
        statute_name = "tax lien sale"
        window_years = "three (3)"
    else:
        statute_cite = "C.R.S. § 38-38-111"
        statute_name = "public trustee foreclosure sale"
        window_years = "five (5)"

    # Court caption
    _add_styled_paragraph(
        doc,
        f"DISTRICT COURT, {county.upper()} COUNTY, STATE OF COLORADO",
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )

    _add_styled_paragraph(
        doc, "____________________________________________",
        size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
        color=MID_GRAY,
    )

    # Case style block
    style_table = doc.add_table(rows=3, cols=2)
    style_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Petitioner
    p1 = style_table.cell(0, 0).paragraphs[0]
    r1 = p1.add_run(f"In the Matter of Surplus Funds\nCase No. {case_number}")
    r1.font.size = Pt(11)
    r1.font.name = "Times New Roman"

    # v.
    p2 = style_table.cell(1, 0).paragraphs[0]
    r2 = p2.add_run(f"\nPetitioner: {owner}")
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.name = "Times New Roman"

    # Case number on right
    p3 = style_table.cell(0, 1).paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run(f"Case No.: {case_number}")
    r3.font.size = Pt(10)
    r3.font.name = "Times New Roman"

    p4 = style_table.cell(1, 1).paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r4 = p4.add_run(f"Division: ______")
    r4.font.size = Pt(10)
    r4.font.name = "Times New Roman"

    _add_styled_paragraph(
        doc, "____________________________________________",
        size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
        color=MID_GRAY,
    )

    # ---- MOTION TITLE ----
    _add_styled_paragraph(
        doc,
        "MOTION FOR DISBURSEMENT OF SURPLUS FUNDS",
        size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4, font_name="Times New Roman",
    )

    _add_styled_paragraph(
        doc,
        f"Pursuant to {statute_cite}",
        size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16, font_name="Times New Roman", color=MID_GRAY,
    )

    # ---- MOTION BODY ----
    # Paragraph 1: Introduction
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Inches(0.5)

    body_font = "Times New Roman"
    body_size = Pt(11)

    run = p.add_run(
        f"COMES NOW the Petitioner, {owner}, by and through undersigned counsel, "
        f"and respectfully moves this Court for an Order directing the "
        f"{county} County Public Trustee to disburse surplus funds held in "
        f"connection with the {statute_name} of the property located at "
        f"{property_address}, and in support thereof states as follows:"
    )
    run.font.size = body_size
    run.font.name = body_font

    # Numbered paragraphs
    numbered_paras = [
        (
            f"On or about {sale_date}, a {statute_name} was conducted on the "
            f"real property located at {property_address}, {county} County, Colorado "
            f"(the \"Property\"), under Case No. {case_number}."
        ),
        (
            f"The {statute_name} generated proceeds in excess of the amounts owed "
            f"on the {lien_type} and associated costs. The estimated surplus funds "
            f"available for disbursement total approximately {surplus_str}."
        ),
        (
            f"Pursuant to {statute_cite}, surplus funds remaining after satisfaction "
            f"of the obligation secured by the deed of trust, including costs and "
            f"expenses of the sale, shall be paid to the person or persons legally "
            f"entitled thereto."
        ),
        (
            f"Petitioner, {owner}, was the owner of record of the Property at the "
            f"time of the {statute_name} and is the person legally entitled to "
            f"receive the surplus funds pursuant to {statute_cite}."
        ),
        (
            f"Pursuant to {statute_cite}, a claim for surplus funds must be filed "
            f"within {window_years} years of the date of the sale. This Motion is "
            f"timely filed within the statutory period."
        ),
        (
            "No other party has a superior claim to the surplus funds. Petitioner "
            "is unaware of any junior liens, judgments, or encumbrances that would "
            "take priority over Petitioner's claim to the surplus. [ATTORNEY: VERIFY "
            "JUNIOR LIEN STATUS BEFORE FILING.]"
        ),
    ]

    for i, text in enumerate(numbered_paras, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.5)

        num_run = p.add_run(f"{i}.   ")
        num_run.font.size = body_size
        num_run.font.bold = True
        num_run.font.name = body_font

        text_run = p.add_run(text)
        text_run.font.size = body_size
        text_run.font.name = body_font

    # ---- PRAYER FOR RELIEF ----
    _add_styled_paragraph(doc, "", size=6, space_after=4)  # Spacer

    _add_styled_paragraph(
        doc, "PRAYER FOR RELIEF",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8, font_name="Times New Roman",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(
        f"WHEREFORE, Petitioner respectfully requests that this Court enter "
        f"an Order directing the {county} County Public Trustee to disburse "
        f"the surplus funds held in Case No. {case_number}, in the approximate "
        f"amount of {surplus_str}, to the Petitioner, {owner}, together with "
        f"any accrued interest thereon, and for such other and further relief "
        f"as this Court deems just and proper."
    )
    run.font.size = body_size
    run.font.name = body_font

    # ---- SIGNATURE BLOCK ----
    _add_styled_paragraph(doc, "", size=6, space_after=20)  # Spacer

    sig_lines = [
        "Respectfully submitted,",
        "",
        "",
        "________________________________________",
        "[ATTORNEY NAME], Esq.",
        "[BAR NUMBER]",
        "[FIRM NAME]",
        "[ADDRESS]",
        "[PHONE]  |  [EMAIL]",
        "",
        f"Attorney for Petitioner, {owner}",
    ]

    for line in sig_lines:
        _add_styled_paragraph(
            doc, line,
            size=10 if line.startswith("[") or line.startswith("_") else 11,
            color=MID_GRAY if line.startswith("[") else DARK_GRAY,
            font_name="Times New Roman",
            space_after=2,
        )

    # ---- CERTIFICATE OF SERVICE ----
    _add_styled_paragraph(doc, "", size=6, space_after=8)

    _add_styled_paragraph(
        doc, "CERTIFICATE OF SERVICE",
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8, font_name="Times New Roman",
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(
        "I hereby certify that on this ___ day of ____________, 20__, "
        "a true and correct copy of the foregoing MOTION FOR DISBURSEMENT "
        "OF SURPLUS FUNDS was served upon all parties of record by "
        "[U.S. Mail / Electronic Filing / Hand Delivery]."
    )
    run.font.size = Pt(10)
    run.font.name = body_font

    _add_styled_paragraph(doc, "", size=6, space_after=16)

    cert_sig = [
        "________________________________________",
        "[ATTORNEY NAME]",
    ]
    for line in cert_sig:
        _add_styled_paragraph(
            doc, line,
            size=10, color=MID_GRAY if line.startswith("[") else DARK_GRAY,
            font_name="Times New Roman", space_after=2,
        )


# ============================================================================
# DISCLAIMER FOOTER (bottom of every page via section footer)
# ============================================================================

def _add_footer(doc: Document):
    """Add confidentiality footer to the document."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)

    run = p.add_run(
        "CONFIDENTIAL — ATTORNEY WORK PRODUCT  |  "
        "VeriFuse Legal Intelligence  |  verifuse.tech"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = MID_GRAY
    run.font.name = "Calibri"


# ============================================================================
# MAIN GENERATOR
# ============================================================================

def generate_dossier(data: Dict, output_path: Optional[str] = None) -> str:
    """
    Generate a 3-page Pre-Litigation Evidence Packet as a Word document.

    Args:
        data: Asset dictionary with fields from the pipeline or hunter engine.
              Required: county, case_number, estimated_surplus
              Optional: all other fields (gracefully handled if missing)

        output_path: Where to save the .docx file.
                     Defaults to "VF_DOSSIER_{asset_id}_{timestamp}.docx"

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    # ---- PAGE SETUP ----
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ---- BUILD PAGES ----
    _build_page_1(doc, data)
    _build_page_2(doc, data)
    _build_page_3(doc, data)
    _add_footer(doc)

    # ---- SAVE ----
    if not output_path:
        asset_id = data.get("asset_id", "UNKNOWN")
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = f"VF_DOSSIER_{asset_id}_{timestamp}.docx"

    doc.save(output_path)
    print(f"Dossier generated: {output_path}")
    return output_path


def generate_batch(records: list, output_dir: str = "dossiers") -> list:
    """
    Generate dossiers for a batch of records.

    Args:
        records: List of asset dicts (from hunter engine or pipeline)
        output_dir: Directory to save files

    Returns:
        List of generated file paths
    """
    os.makedirs(output_dir, exist_ok=True)
    paths = []

    for i, record in enumerate(records):
        asset_id = record.get("asset_id", f"ASSET_{i:04d}")
        county = record.get("county", "UNK").replace(" ", "_")
        filename = f"VF_DOSSIER_{county}_{asset_id}.docx"
        filepath = os.path.join(output_dir, filename)

        try:
            generate_dossier(record, filepath)
            paths.append(filepath)
        except Exception as e:
            print(f"ERROR generating dossier for {asset_id}: {e}")

    print(f"\nBatch complete: {len(paths)}/{len(records)} dossiers generated in {output_dir}/")
    return paths


# ============================================================================
# SAMPLE DATA & CLI
# ============================================================================

SAMPLE_ASSET = {
    "asset_id": "Denver_FORECLOSURE_SURPLUS_a3f8c912",
    "county": "Denver",
    "state": "CO",
    "jurisdiction": "CO_Denver",
    "asset_type": "FORECLOSURE_SURPLUS",
    "case_number": "2025CV030892",
    "property_address": "1847 WILLIAMS ST, DENVER, CO 80218",
    "owner_of_record": "Margaret A. Rodriguez",
    "estimated_surplus": 87432.19,
    "total_indebtedness": 312567.81,
    "overbid_amount": 400000.00,
    "sale_date": "2025-11-14",
    "lien_type": "Deed of Trust",
    "recorder_link": "https://denvergov.org/recorder/search?query=Margaret+A.+Rodriguez",
    "redemption_date": "2026-05-14",
    "statute_years": 5,
    "data_grade": "GOLD",
    "source_file": "hunter:denver:public_trustee_scrape",
    "created_at": "2026-01-15T08:30:00Z",
    "updated_at": "2026-02-01T14:22:00Z",

    # Intelligence fields from hunter engine
    "_mailing_address": "4521 E COLFAX AVE APT 302, DENVER, CO 80220",
    "_is_absentee": True,
    "_absentee_reason": "address_mismatch",
    "_classification": "PRIME",
    "_litigation_quality": "A",
    "_junior_lien_count": 0,
    "_estimated_fee": 28853.62,

    # Optional
    "satellite_map_url": (
        "https://maps.googleapis.com/maps/api/staticmap?"
        "center=1847+Williams+St+Denver+CO&zoom=17&size=600x300&maptype=satellite"
    ),
    "document_link": "",
}


if __name__ == "__main__":
    print("=" * 70)
    print("VERIFUSE DOSSIER GENERATOR — Sample Output")
    print("=" * 70)

    path = generate_dossier(SAMPLE_ASSET)
    print(f"\nSample dossier saved to: {path}")
    print("\nTo generate for real assets:")
    print("  from verifuse.attorney.dossier_generator import generate_dossier")
    print("  generate_dossier(asset_dict, 'output.docx')")
