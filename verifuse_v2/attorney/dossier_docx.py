"""
VeriFuse V2 — DOCX Dossier Generator
=====================================
Generates a 3-page Word Document "Forensic Report" for a lead.

PAGE 1: Executive Intelligence Summary
PAGE 2: Chain of Title Audit
PAGE 3: Ready-to-File Motion for Disbursement

Ported from V1 verifuse/attorney/dossier_generator.py with changes:
  - Accepts lead_id, queries `leads` table directly
  - Uses V2 column names (owner_name, total_debt, surplus_amount)
  - Pre-transfer period is "Unregulated (no statutory cap)"
  - Outputs to verifuse_v2/data/dossiers/

Usage:
    python -m verifuse_v2.attorney.dossier_docx --lead-id <LEAD_ID>
    python -m verifuse_v2.attorney.dossier_docx --lead-id <LEAD_ID> --db /path/to/db
"""

from __future__ import annotations

import os
import sqlite3
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Color palette ──────────────────────────────────────────────────
BLACK = RGBColor(0x0A, 0x19, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_GREEN = RGBColor(0x00, 0x80, 0x3E)
ACCENT_RED = RGBColor(0xCC, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)

DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "dossiers"


# ── Helpers ────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_paragraph(doc, text: str, size: int = 11,
                          bold: bool = False, italic: bool = False,
                          color: RGBColor = DARK_GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_after: int = 6, space_before: int = 0,
                          font_name: str = "Calibri"):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font_name
    return p


def _format_currency(amount) -> str:
    if amount is None:
        return "NOT AVAILABLE"
    try:
        val = float(amount)
        if val < 0:
            return f"-${abs(val):,.2f}"
        return f"${val:,.2f}"
    except (ValueError, TypeError):
        return str(amount)


def _compute_statute_status(sale_date_str: Optional[str], statute_years: int = 5) -> dict:
    if not sale_date_str:
        return {
            "label": "INSUFFICIENT DATA",
            "days_since_sale": None,
            "days_remaining": None,
            "fee_cap": None,
            "fee_cap_label": "Unknown",
        }
    try:
        sale_date = datetime.strptime(str(sale_date_str)[:10], "%Y-%m-%d")
    except ValueError:
        return {
            "label": "DATE PARSE ERROR",
            "days_since_sale": None,
            "days_remaining": None,
            "fee_cap": None,
            "fee_cap_label": "Unknown",
        }

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    days_since = (now - sale_date).days
    total_window = statute_years * 365
    days_remaining = total_window - days_since

    if days_since <= 180:
        label = "DATA ACCESS ONLY — C.R.S. § 38-38-111(2.5)(c)"
        fee_cap = None
        fee_cap_label = "Compensation agreements void during holding period"
    elif days_since <= total_window:
        label = "ESCROW ENDED — Holding period expired"
        fee_cap = None
        fee_cap_label = "Consult C.R.S. § 38-38-111 for applicable rules"
    else:
        label = "EXPIRED — Funds may have escheated"
        fee_cap = None
        fee_cap_label = "C.R.S. § 38-13-101 (unclaimed property)"

    return {
        "label": label,
        "days_since_sale": days_since,
        "days_remaining": max(days_remaining, 0),
        "fee_cap": fee_cap,
        "fee_cap_label": fee_cap_label,
    }


def _fetch_lead(db_path: str, lead_id: str) -> dict:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM leads WHERE id = ?", [lead_id]).fetchone()
    conn.close()
    if not row:
        raise ValueError(f"Lead {lead_id} not found in database")
    return dict(row)


# ── PAGE 1: Executive Intelligence Summary ─────────────────────────

def _build_page_1(doc: Document, data: dict):
    # Header bar
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    _set_cell_shading(header_cell, "0A192F")
    header_cell.width = Inches(7.0)

    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(12)
    hp.paragraph_format.space_after = Pt(12)

    run = hp.add_run("VERIFUSE FORENSIC REPORT")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = "Calibri"

    run2 = hp.add_run("  |  CONFIDENTIAL")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x64, 0xFF, 0xDA)
    run2.font.name = "Calibri"

    # Metadata
    generated = datetime.now(timezone.utc).strftime("%B %d, %Y at %H:%M UTC")
    _add_styled_paragraph(
        doc,
        f"Generated: {generated}  |  Lead: {data.get('id', 'N/A')}  |  "
        f"County: {data.get('county', 'N/A')}, CO",
        size=9, color=MID_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=16,
    )

    # Section title
    _add_styled_paragraph(
        doc, "EXECUTIVE INTELLIGENCE SUMMARY",
        size=13, bold=True, color=BLACK, space_before=4, space_after=8,
    )

    surplus = (data.get("surplus_amount") or data.get("estimated_surplus") or 0)
    statute = _compute_statute_status(data.get("sale_date"))

    # Metrics table
    metrics = doc.add_table(rows=2, cols=2)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.style = "Table Grid"

    for row in metrics.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.8)

    # Row 1: Liquidity
    label_cell = metrics.cell(0, 0)
    _set_cell_shading(label_cell, "0A192F")
    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after = Pt(8)
    lr = lp.add_run("LIQUIDITY\nAVAILABLE")
    lr.font.size = Pt(10)
    lr.font.bold = True
    lr.font.color.rgb = WHITE
    lr.font.name = "Calibri"

    value_cell = metrics.cell(0, 1)
    vp = value_cell.paragraphs[0]
    vp.paragraph_format.space_before = Pt(6)
    vp.paragraph_format.space_after = Pt(6)
    vr = vp.add_run(_format_currency(surplus))
    vr.font.size = Pt(22)
    vr.font.bold = True
    vr.font.color.rgb = ACCENT_GREEN if surplus and surplus > 0 else ACCENT_RED
    vr.font.name = "Calibri"

    # Fee cap removed per C.R.S. § 38-38-111 compliance

    # Row 2: Statute Window
    label_cell_2 = metrics.cell(1, 0)
    _set_cell_shading(label_cell_2, "0A192F")
    lp2 = label_cell_2.paragraphs[0]
    lp2.paragraph_format.space_before = Pt(8)
    lp2.paragraph_format.space_after = Pt(8)
    lr2 = lp2.add_run("STATUTE\nWINDOW")
    lr2.font.size = Pt(10)
    lr2.font.bold = True
    lr2.font.color.rgb = WHITE
    lr2.font.name = "Calibri"

    value_cell_2 = metrics.cell(1, 1)
    vp2 = value_cell_2.paragraphs[0]
    vp2.paragraph_format.space_before = Pt(6)
    vp2.paragraph_format.space_after = Pt(6)
    vr2 = vp2.add_run(statute["label"])
    vr2.font.size = Pt(12)
    vr2.font.bold = True
    vr2.font.name = "Calibri"

    if "PRE-TRANSFER" in statute["label"]:
        vr2.font.color.rgb = ACCENT_GREEN
    elif "FINDER" in statute["label"]:
        vr2.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
    else:
        vr2.font.color.rgb = ACCENT_RED

    if statute["days_remaining"] is not None:
        dr_text = f"\n{statute['days_remaining']} days remaining  |  Fee cap: {statute['fee_cap_label']}"
        dr_run = vp2.add_run(dr_text)
        dr_run.font.size = Pt(9)
        dr_run.font.color.rgb = MID_GRAY
        dr_run.font.name = "Calibri"

    # Property details
    _add_styled_paragraph(doc, "", size=6, space_after=2)

    details_table = doc.add_table(rows=4, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.style = "Table Grid"

    detail_rows = [
        ("Property Address", data.get("property_address") or "NOT AVAILABLE"),
        ("Owner of Record", data.get("owner_name") or "NOT AVAILABLE"),
        ("Case / File Number", data.get("case_number") or "NOT AVAILABLE"),
        ("Data Grade", data.get("data_grade") or "N/A"),
    ]
    for i, (label, value) in enumerate(detail_rows):
        lc = details_table.cell(i, 0)
        vc = details_table.cell(i, 1)
        lc.width = Inches(2.2)
        vc.width = Inches(4.8)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(9)
        lr.font.bold = True
        lr.font.color.rgb = DARK_GRAY
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(10)
        vr.font.color.rgb = DARK_GRAY
        vr.font.name = "Calibri"


# ── PAGE 2: Chain of Title Audit ───────────────────────────────────

def _build_page_2(doc: Document, data: dict):
    doc.add_page_break()

    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    _set_cell_shading(header_cell, "0A192F")
    header_cell.width = Inches(7.0)

    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(8)
    hp.paragraph_format.space_after = Pt(8)
    run = hp.add_run("CHAIN OF TITLE AUDIT")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = "Calibri"

    _add_styled_paragraph(
        doc, "Foreclosure / Sale Event",
        size=12, bold=True, color=BLACK, space_before=14, space_after=6,
    )

    event_data = [
        ("Foreclosure / Sale Date", data.get("sale_date") or "NOT AVAILABLE"),
        ("County", f"{data.get('county', 'N/A')}, CO"),
        ("Case Number", data.get("case_number") or "NOT AVAILABLE"),
        ("Claim Deadline", data.get("claim_deadline") or "N/A"),
    ]

    event_table = doc.add_table(rows=len(event_data), cols=2)
    event_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    event_table.style = "Table Grid"

    for i, (label, value) in enumerate(event_data):
        lc = event_table.cell(i, 0)
        vc = event_table.cell(i, 1)
        lc.width = Inches(2.5)
        vc.width = Inches(4.5)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(10)
        lr.font.bold = True
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(10)
        vr.font.name = "Calibri"

    # THE MATH
    _add_styled_paragraph(
        doc, "THE MATH — Surplus Calculation",
        size=12, bold=True, color=BLACK, space_before=20, space_after=4,
    )

    bid = data.get("winning_bid") or data.get("overbid_amount")
    debt = data.get("total_debt")
    surplus = data.get("surplus_amount") or data.get("estimated_surplus")

    if surplus is None and bid is not None and debt is not None:
        try:
            surplus = float(bid) - float(debt)
        except (ValueError, TypeError):
            pass

    math_table = doc.add_table(rows=4, cols=2)
    math_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    math_table.style = "Table Grid"

    math_rows = [
        ("Winning Bid / Sale Price", _format_currency(bid)),
        ("Total Debt / Judgment", f"– {_format_currency(debt)}"),
        ("", ""),
        ("SURPLUS (Liquidity Available)", _format_currency(surplus)),
    ]

    for i, (label, value) in enumerate(math_rows):
        lc = math_table.cell(i, 0)
        vc = math_table.cell(i, 1)
        lc.width = Inches(3.5)
        vc.width = Inches(3.5)

        if i == 2:
            _set_cell_shading(lc, "0A192F")
            _set_cell_shading(vc, "0A192F")
            lc.paragraphs[0].paragraph_format.space_before = Pt(1)
            lc.paragraphs[0].paragraph_format.space_after = Pt(1)
            vc.paragraphs[0].paragraph_format.space_before = Pt(1)
            vc.paragraphs[0].paragraph_format.space_after = Pt(1)
            continue

        if i == 3:
            _set_cell_shading(lc, "E8F5E9")
            _set_cell_shading(vc, "E8F5E9")

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.font.size = Pt(11)
        lr.font.bold = (i == 3)
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vr = vp.add_run(value)
        vr.font.size = Pt(14 if i == 3 else 11)
        vr.font.bold = (i == 3)
        vr.font.name = "Calibri"
        if i == 3:
            vr.font.color.rgb = ACCENT_GREEN if surplus and surplus > 0 else ACCENT_RED

    # Data Provenance
    _add_styled_paragraph(
        doc, "DATA PROVENANCE",
        size=11, bold=True, color=BLACK, space_before=16, space_after=4,
    )

    prov_data = [
        ("Data Source", f"Public records — {data.get('county', 'N/A')} County, CO"),
        ("Last Updated", str(data.get("updated_at", "Unknown"))[:10]),
        ("Data Grade", data.get("data_grade") or "N/A"),
        ("Source", data.get("source_name") or "VeriFuse automated collection"),
    ]

    prov_table = doc.add_table(rows=len(prov_data), cols=2)
    prov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prov_table.style = "Table Grid"

    for i, (label, value) in enumerate(prov_data):
        lc = prov_table.cell(i, 0)
        vc = prov_table.cell(i, 1)
        lc.width = Inches(2.0)
        vc.width = Inches(5.0)
        _set_cell_shading(lc, "F2F2F2")

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(8)
        lr.font.bold = True
        lr.font.name = "Calibri"

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(8)
        vr.font.name = "Calibri"


# ── PAGE 3: Ready-to-File Motion ──────────────────────────────────

def _build_page_3(doc: Document, data: dict):
    doc.add_page_break()

    county = data.get("county", "[COUNTY]")
    case_number = data.get("case_number", "[CASE NUMBER]")
    owner = data.get("owner_name", "[OWNER NAME]")
    property_address = data.get("property_address", "[PROPERTY ADDRESS]")
    surplus = data.get("surplus_amount") or data.get("estimated_surplus")
    surplus_str = _format_currency(surplus)
    sale_date = data.get("sale_date", "[SALE DATE]")

    statute_cite = "C.R.S. § 38-38-111"
    statute_name = "public trustee foreclosure sale"
    window_years = "five (5)"

    _add_styled_paragraph(
        doc,
        f"DISTRICT COURT, {county.upper()} COUNTY, STATE OF COLORADO",
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )

    _add_styled_paragraph(
        doc, "____________________________________________",
        size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
        color=MID_GRAY,
    )

    # Case style
    style_table = doc.add_table(rows=2, cols=2)
    style_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    p1 = style_table.cell(0, 0).paragraphs[0]
    r1 = p1.add_run(f"In the Matter of Surplus Funds\nCase No. {case_number}")
    r1.font.size = Pt(11)
    r1.font.name = "Times New Roman"

    p2 = style_table.cell(1, 0).paragraphs[0]
    r2 = p2.add_run(f"\nPetitioner: {owner}")
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.name = "Times New Roman"

    p3 = style_table.cell(0, 1).paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run(f"Case No.: {case_number}")
    r3.font.size = Pt(10)
    r3.font.name = "Times New Roman"

    p4 = style_table.cell(1, 1).paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r4 = p4.add_run("Division: ______")
    r4.font.size = Pt(10)
    r4.font.name = "Times New Roman"

    _add_styled_paragraph(
        doc, "____________________________________________",
        size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
        color=MID_GRAY,
    )

    _add_styled_paragraph(
        doc,
        "MOTION FOR DISBURSEMENT OF SURPLUS FUNDS",
        size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4, font_name="Times New Roman",
    )

    _add_styled_paragraph(
        doc,
        f"Pursuant to {statute_cite}",
        size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16, font_name="Times New Roman", color=MID_GRAY,
    )

    body_font = "Times New Roman"
    body_size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(
        f"COMES NOW the Petitioner, {owner}, by and through undersigned counsel, "
        f"and respectfully moves this Court for an Order directing the "
        f"{county} County Public Trustee to disburse surplus funds held in "
        f"connection with the {statute_name} of the property located at "
        f"{property_address}, and in support thereof states as follows:"
    )
    run.font.size = body_size
    run.font.name = body_font

    numbered_paras = [
        (
            f"On or about {sale_date}, a {statute_name} was conducted on the "
            f"real property located at {property_address}, {county} County, Colorado "
            f'(the "Property"), under Case No. {case_number}.'
        ),
        (
            f"The {statute_name} generated proceeds in excess of the amounts owed "
            f"on the deed of trust and associated costs. The estimated surplus funds "
            f"available for disbursement total approximately {surplus_str}."
        ),
        (
            f"Pursuant to {statute_cite}, surplus funds remaining after satisfaction "
            f"of the obligation secured by the deed of trust, including costs and "
            f"expenses of the sale, shall be paid to the person or persons legally "
            f"entitled thereto."
        ),
        (
            f"Petitioner, {owner}, was the owner of record of the Property at the "
            f"time of the {statute_name} and is the person legally entitled to "
            f"receive the surplus funds pursuant to {statute_cite}."
        ),
        (
            f"Pursuant to {statute_cite}, a claim for surplus funds must be filed "
            f"within {window_years} years of the date of the sale. This Motion is "
            f"timely filed within the statutory period."
        ),
        (
            "No other party has a superior claim to the surplus funds. Petitioner "
            "is unaware of any junior liens, judgments, or encumbrances that would "
            "take priority over Petitioner's claim to the surplus. [ATTORNEY: VERIFY "
            "JUNIOR LIEN STATUS BEFORE FILING.]"
        ),
    ]

    for i, text in enumerate(numbered_paras, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.5)

        num_run = p.add_run(f"{i}.   ")
        num_run.font.size = body_size
        num_run.font.bold = True
        num_run.font.name = body_font

        text_run = p.add_run(text)
        text_run.font.size = body_size
        text_run.font.name = body_font

    # Prayer for relief
    _add_styled_paragraph(doc, "", size=6, space_after=4)
    _add_styled_paragraph(
        doc, "PRAYER FOR RELIEF",
        size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8, font_name="Times New Roman",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(
        f"WHEREFORE, Petitioner respectfully requests that this Court enter "
        f"an Order directing the {county} County Public Trustee to disburse "
        f"the surplus funds held in Case No. {case_number}, in the approximate "
        f"amount of {surplus_str}, to the Petitioner, {owner}, together with "
        f"any accrued interest thereon, and for such other and further relief "
        f"as this Court deems just and proper."
    )
    run.font.size = body_size
    run.font.name = body_font

    # Signature block
    _add_styled_paragraph(doc, "", size=6, space_after=20)

    sig_lines = [
        "Respectfully submitted,",
        "", "",
        "________________________________________",
        "[ATTORNEY NAME], Esq.",
        "[BAR NUMBER]",
        "[FIRM NAME]",
        "[ADDRESS]",
        "[PHONE]  |  [EMAIL]",
        "",
        f"Attorney for Petitioner, {owner}",
    ]

    for line in sig_lines:
        _add_styled_paragraph(
            doc, line,
            size=10 if line.startswith("[") or line.startswith("_") else 11,
            color=MID_GRAY if line.startswith("[") else DARK_GRAY,
            font_name="Times New Roman", space_after=2,
        )


# ── Footer ─────────────────────────────────────────────────────────

def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)

    run = p.add_run(
        "CONFIDENTIAL — ATTORNEY WORK PRODUCT  |  "
        "VeriFuse Legal Intelligence  |  verifuse.tech"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = MID_GRAY
    run.font.name = "Calibri"


# ── Main Generator ─────────────────────────────────────────────────

def generate_dossier(db_path: str, lead_id: str, output_dir: str = None) -> str:
    """Generate a 3-page DOCX dossier for a lead.

    Args:
        db_path: Path to the SQLite database.
        lead_id: The lead ID to generate the dossier for.
        output_dir: Directory to save the .docx file. Defaults to data/dossiers/.

    Returns:
        Path to the generated .docx file.
    """
    data = _fetch_lead(db_path, lead_id)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    _build_page_1(doc, data)
    _build_page_2(doc, data)
    _build_page_3(doc, data)
    _add_footer(doc)

    out_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    county = (data.get("county") or "UNK").replace(" ", "_")
    short_id = str(lead_id)[:12]
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"VF_DOSSIER_{county}_{short_id}_{timestamp}.docx"
    filepath = out_dir / filename

    doc.save(str(filepath))
    return str(filepath)


# ── CLI ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generate DOCX dossier for a lead")
    parser.add_argument("--lead-id", required=True, help="Lead ID")
    parser.add_argument("--db", default=os.environ.get("VERIFUSE_DB_PATH"),
                        help="Path to database (default: VERIFUSE_DB_PATH env)")
    parser.add_argument("--output-dir", default=None, help="Output directory")
    args = parser.parse_args()

    if not args.db:
        print("FATAL: --db or VERIFUSE_DB_PATH required")
        sys.exit(1)

    path = generate_dossier(args.db, args.lead_id, args.output_dir)
    print(f"Dossier generated: {path}")
