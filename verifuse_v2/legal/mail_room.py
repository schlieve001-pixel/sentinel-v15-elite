"""
VeriFuse V2 — Mail Room (Rule 7.3 Solicitation Letters)
========================================================
Generates CO Rule 7.3-compliant .docx solicitation letters.

Ported from V1 verifuse/legal/mail_room.py with changes:
  - Queries `leads` table instead of `assets` + `legal_status` join
  - Uses VERIFUSE_DB_PATH env var
  - Attorney info from `users` table
  - No 33%/0.33 fee references
  - Outputs to verifuse_v2/data/letters/

Usage:
    python -m verifuse_v2.legal.mail_room --attorney-id <USER_ID>
"""

from __future__ import annotations

import os
import sqlite3
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

COMPLIANCE_FOOTER = "ADVERTISING MATERIAL - COMPLIANT WITH CO RULE 7.3"

DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "letters"

BLACK = RGBColor(0x0A, 0x19, 0x2F)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)


def _fmt_money(amount) -> str:
    if amount is None:
        return "$[AMOUNT]"
    try:
        return f"${float(amount):,.2f}"
    except (ValueError, TypeError):
        return "$[AMOUNT]"


def _add_paragraph(doc, text, size=11, bold=False, color=None,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                   font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _get_attorney_info(db_path: str, attorney_user_id: str) -> dict:
    """Fetch attorney info from users table."""
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        "SELECT * FROM users WHERE user_id = ?", [attorney_user_id]
    ).fetchone()
    conn.close()
    if not row:
        raise ValueError(f"Attorney user {attorney_user_id} not found")
    u = dict(row)
    return {
        "firm_name": u.get("firm_name") or "[LAW FIRM NAME]",
        "attorney_name": u.get("full_name") or "[ATTORNEY NAME, ESQ.]",
        "bar_number": u.get("bar_number") or "[CO BAR #]",
        "address_line1": "[FIRM ADDRESS LINE 1]",
        "address_line2": "[CITY, STATE ZIP]",
        "phone": "[PHONE NUMBER]",
        "email": u.get("email") or "[EMAIL]",
    }


def generate_letter(db_path: str, lead_id: str, attorney_user_id: str,
                    output_dir: str = None) -> str:
    """Generate a Rule 7.3 compliant solicitation letter for a lead.

    Args:
        db_path: Path to the SQLite database.
        lead_id: The lead ID.
        attorney_user_id: The attorney's user_id from the users table.
        output_dir: Directory to save. Defaults to data/letters/.

    Returns:
        Path to the generated .docx file.
    """
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    lead_row = conn.execute("SELECT * FROM leads WHERE id = ?", [lead_id]).fetchone()
    conn.close()
    if not lead_row:
        raise ValueError(f"Lead {lead_id} not found")
    lead = dict(lead_row)

    atty = _get_attorney_info(db_path, attorney_user_id)

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Letterhead
    _add_paragraph(doc, atty["firm_name"], size=14, bold=True, color=BLACK,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _add_paragraph(doc, atty.get("address_line1", ""), size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _add_paragraph(doc, atty.get("address_line2", ""), size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    contact_line = f"Tel: {atty.get('phone', '')}  |  {atty.get('email', '')}"
    _add_paragraph(doc, contact_line, size=9, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    _add_paragraph(doc, "_" * 72, size=8, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    today = datetime.now(timezone.utc).strftime("%B %d, %Y")
    _add_paragraph(doc, today, size=11, color=DARK_GRAY, space_after=12)

    owner = lead.get("owner_name") or "[PROPERTY OWNER]"
    prop_addr = lead.get("property_address") or "[PROPERTY ADDRESS]"

    _add_paragraph(doc, owner, size=11, bold=True, color=DARK_GRAY, space_after=1)
    _add_paragraph(doc, prop_addr, size=11, color=DARK_GRAY, space_after=12)

    surplus = lead.get("surplus_amount") or lead.get("estimated_surplus")
    case_num = lead.get("case_number") or "[CASE NUMBER]"
    county = lead.get("county") or "[COUNTY]"

    re_line = f"RE: Unclaimed Funds of {_fmt_money(surplus)} - Case #{case_num}"
    _add_paragraph(doc, re_line, size=12, bold=True, color=BLACK, space_after=12)

    _add_paragraph(doc, f"Dear {owner},", size=11, color=DARK_GRAY, space_after=8)

    body1 = (
        f"Our office has identified unclaimed surplus funds in the amount of "
        f"{_fmt_money(surplus)} that may belong to you. These funds resulted "
        f"from a foreclosure sale involving real property located at "
        f"{prop_addr}, in {county} County, Colorado."
    )
    _add_paragraph(doc, body1, size=11, color=DARK_GRAY, space_after=8)

    statute = "C.R.S. Section 38-38-111"
    body2 = (
        f"Under {statute}, these funds are held by the {county} County "
        f"Public Trustee and are available for claim by the rightful owner. "
        f"However, there is a statutory deadline to file your claim. "
        f"After this deadline, the funds may be transferred to the Colorado "
        f"State Treasury, making recovery significantly more difficult."
    )
    _add_paragraph(doc, body2, size=11, color=DARK_GRAY, space_after=8)

    body3 = (
        f"Our firm specializes in recovering surplus funds from foreclosure "
        f"sales in Colorado. We handle the entire legal process, including "
        f"filing the necessary petition with the court, locating and "
        f"assembling all required documentation, and representing your "
        f"interests through to the disbursement of funds."
    )
    _add_paragraph(doc, body3, size=11, color=DARK_GRAY, space_after=8)

    body4 = (
        "There is no upfront cost to you. Our fee is contingent upon "
        "successful recovery of your funds. If we do not recover any "
        "money, you owe us nothing."
    )
    _add_paragraph(doc, body4, size=11, bold=True, color=DARK_GRAY, space_after=8)

    body5 = (
        "To discuss this matter, please contact our office at your earliest "
        f"convenience by calling {atty.get('phone', '[PHONE]')} or by "
        f"emailing {atty.get('email', '[EMAIL]')}. Time is of the essence."
    )
    _add_paragraph(doc, body5, size=11, color=DARK_GRAY, space_after=16)

    _add_paragraph(doc, "Sincerely,", size=11, color=DARK_GRAY, space_after=24)
    _add_paragraph(doc, atty["attorney_name"], size=11, bold=True,
                   color=DARK_GRAY, space_after=1)
    _add_paragraph(doc, atty["firm_name"], size=10, color=MID_GRAY, space_after=1)
    bar = atty.get("bar_number", "")
    if bar:
        _add_paragraph(doc, f"Colorado Bar #{bar}", size=9, color=MID_GRAY,
                       space_after=1)

    # Compliance footer (MANDATORY)
    _add_paragraph(doc, "", size=6, space_after=20)
    _add_paragraph(doc, "_" * 72, size=8, color=MID_GRAY,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _add_paragraph(
        doc, COMPLIANCE_FOOTER,
        size=9, bold=True, color=RGBColor(0xCC, 0x00, 0x00),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_paragraph(
        doc,
        "This letter is being sent to you because public records indicate "
        "you may have a legal right to surplus funds. You are under no "
        "obligation to respond. You may wish to consult with another "
        "attorney of your choosing.",
        size=8, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )

    out_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    county_clean = county.replace(" ", "_")
    short_id = str(lead_id)[:12]
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d")
    filename = f"VF_LETTER_{county_clean}_{short_id}_{timestamp}.docx"
    filepath = out_dir / filename

    doc.save(str(filepath))
    return str(filepath)


def get_qualifying_leads(db_path: str, min_surplus: float = 25000) -> list:
    """Get leads qualifying for attorney solicitation letters.

    Returns list of lead dicts with surplus >= min_surplus, valid owner, active window.
    """
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("""
        SELECT * FROM leads
        WHERE data_grade IN ('GOLD', 'SILVER')
          AND COALESCE(surplus_amount, estimated_surplus, 0) >= ?
          AND owner_name IS NOT NULL AND owner_name != ''
          AND owner_name != 'Unknown'
          AND attorney_packet_ready = 1
        ORDER BY COALESCE(surplus_amount, estimated_surplus, 0) DESC
    """, [min_surplus]).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def run_mail_room(db_path: str, attorney_user_id: str,
                  output_dir: str = None) -> list:
    """Batch generate solicitation letters for all qualifying leads.

    Args:
        db_path: Path to the SQLite database.
        attorney_user_id: Attorney's user_id from users table.
        output_dir: Directory for generated letters.

    Returns:
        List of generated file paths.
    """
    leads = get_qualifying_leads(db_path)
    if not leads:
        print("No qualifying leads found.")
        return []

    print(f"Generating letters for {len(leads)} leads...")
    paths = []
    for lead in leads:
        try:
            path = generate_letter(db_path, lead["id"], attorney_user_id, output_dir)
            paths.append(path)
            surplus = lead.get("surplus_amount") or lead.get("estimated_surplus") or 0
            owner = (lead.get("owner_name") or "Unknown")[:30]
            print(f"  {lead.get('county', 'UNK'):12s} | {_fmt_money(surplus):>14s} | {owner}")
        except Exception as e:
            print(f"  [ERR] {lead.get('id', '?')}: {e}")

    print(f"\nMail room complete: {len(paths)}/{len(leads)} letters generated")
    return paths


if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="Generate Rule 7.3 solicitation letters")
    parser.add_argument("--attorney-id", required=True, help="Attorney user_id")
    parser.add_argument("--db", default=os.environ.get("VERIFUSE_DB_PATH"),
                        help="Path to database")
    parser.add_argument("--output-dir", default=None, help="Output directory")
    args = parser.parse_args()

    if not args.db:
        print("FATAL: --db or VERIFUSE_DB_PATH required")
        sys.exit(1)

    paths = run_mail_room(args.db, args.attorney_id, args.output_dir)
    print(f"Generated {len(paths)} letters")
